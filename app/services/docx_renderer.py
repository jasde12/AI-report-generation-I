from __future__ import annotations

import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Length, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxcompose.composer import Composer
from app.config import get_settings
from app.models.schemas import (
    EvidenceReference,
    EvidenceText,
    ExtractedSource,
    ExtractedTable,
    NormalizedDocument,
    ReportJSON,
    ReportSection,
    ReportTable,
)
from app.services.analysis_chart_generator import AnalysisChartGenerator
from app.services.template_registry import get_template_definition, get_template_docx_path


class DocxRenderer:
    WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    LIST_STYLE = "List Paragraph"
    BULLET_STYLE = "List Bullet"
    HEADING_STYLE = "Heading 1"
    ATTACHMENT_TITLE_STYLE = "Heading 2"
    BODY_PARAGRAPH_LINE_SPACING = Pt(20)
    BODY_FONT_SIZE = Pt(14)
    BODY_FONT_NAME = "微軟正黑體"
    TOC_TITLE_FONT_SIZE = Pt(16)
    TOC_ENTRY_FONT_SIZE = Pt(10)
    COVER_MINISTRY_FONT_SIZE = Pt(26)
    COVER_PLAN_FONT_SIZE = Pt(19)
    COVER_REPORT_TITLE_FONT_SIZE = Pt(16)
    DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE = Pt(10)
    DIAGNOSIS_SUMMARY_HEADER_FILL = "154F67"
    DIAGNOSIS_SUMMARY_HIGHLIGHT_FILL = "FFF200"
    RECOMMENDATION_INTRO_FONT_SIZE = Pt(14)
    RECOMMENDATION_BODY_FONT_SIZE = Pt(12)
    RECOMMENDATION_INTRO_LEFT_INDENT = Pt(25.5)
    RECOMMENDATION_INTRO_FIRST_LINE_INDENT = Pt(16.95)
    RECOMMENDATION_INTRO_SPACE_AFTER = Pt(8)
    DIAGNOSIS_IMAGE_WIDTH = Cm(11.77)
    DIAGNOSIS_BODY_LEFT_INDENT = Pt(25.5)
    DIAGNOSIS_BODY_FIRST_LINE_INDENT = Pt(16.95)
    DIAGNOSIS_BODY_SPACE_AFTER = Pt(8)
    ANALYSIS_CHART_WIDTH = Cm(15.32)
    INTRODUCTION_LEFT_INDENT = Pt(25.5)
    INTRODUCTION_FIRST_LINE_INDENT = Pt(27.3)
    RECOMMENDATION_SECTION_LEFT_INDENT = Pt(42.55)
    RECOMMENDATION_SECTION_FIRST_LINE_INDENT = Pt(-28.35)
    RECOMMENDATION_ITEM_LEFT_INDENT = Pt(56.6)
    RECOMMENDATION_ITEM_FIRST_LINE_INDENT = Pt(-28.3)
    RECOMMENDATION_FORMULA_LEFT_INDENT = Pt(56.4)
    RECOMMENDATION_FORMULA_FIRST_LINE_INDENT = Pt(-6.85)

    def __init__(self, output_dir: Path | None = None) -> None:
        self.settings = get_settings()
        self.output_dir = output_dir or self.settings.output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.chart_generator = AnalysisChartGenerator(self.output_dir)

    def render(
        self,
        project_name: str,
        template_id: str,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> Path:
        template_definition = get_template_definition(template_id)
        template_path = get_template_docx_path(template_id)

        if template_path is not None:
            document = Document(template_path)
            self._render_d_template_document(
                document=document,
                project_name=project_name,
                report=report,
                normalized_document=normalized_document,
            )
        else:
            document = Document()
            self._render_generic_document(
                document=document,
                project_name=project_name,
                report=report,
                template_name=template_definition.name,
                normalized_document=normalized_document,
            )

        self._apply_document_font(document)
        self._apply_cover_page_format(document)
        self._apply_recommendation_chapter_font(document)
        self._apply_table_of_contents_format(document)
        self._apply_forced_layout_overrides(document)
        if template_path is not None:
            document = self._reload_template_document(document)
        output_name = self._build_output_name(
            self._resolve_output_label(project_name, normalized_document)
        )
        output_path = self.output_dir / output_name
        document.save(output_path)
        return output_path

    def _reload_template_document(self, document: Document) -> Document:
        temp_output_path = self.output_dir / f"._render_tmp_{uuid4().hex}.docx"
        try:
            document.save(temp_output_path)
            reloaded_document = Document(temp_output_path)
            self._apply_document_font(reloaded_document)
            self._apply_cover_page_format(reloaded_document)
            self._apply_recommendation_chapter_font(reloaded_document)
            self._apply_table_of_contents_format(reloaded_document)
            self._apply_forced_layout_overrides(reloaded_document)
            return reloaded_document
        finally:
            temp_output_path.unlink(missing_ok=True)

    def repair_existing_report(
        self,
        input_path: Path,
        output_path: Path | None = None,
    ) -> Path:
        document = Document(input_path)
        self._apply_document_font(document)
        self._apply_cover_page_format(document)
        self._apply_recommendation_chapter_font(document)
        self._apply_table_of_contents_format(document)
        self._apply_forced_layout_overrides(document)
        document = self._reload_template_document(document)

        resolved_output_path = output_path or input_path.with_name(f"{input_path.stem}_repaired.docx")
        document.save(resolved_output_path)
        return resolved_output_path

    def _render_d_template_document(
        self,
        document: Document,
        project_name: str,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> None:
        hospital_name = self._infer_hospital_name(project_name, normalized_document)
        diagnosis_hospital_name = self._with_ministry_prefix(hospital_name)
        report_date = self._format_roc_date()

        self._populate_cover_page(document, hospital_name=hospital_name, report_date=report_date)
        self._populate_table_of_contents(document, hospital_name=diagnosis_hospital_name)

        introduction_heading = self._find_heading_paragraph(document, lambda text: text == "壹、前言")
        diagnosis_heading = self._find_heading_paragraph(
            document,
            lambda text: text.startswith("貳、") and text.endswith("設備韌性備援診斷"),
        )
        recommendation_heading = self._find_heading_paragraph(
            document,
            lambda text: text == "參、孤島效應下備援能力及建議",
        )
        attachment_heading = self._find_heading_paragraph(document, lambda text: text == "肆、附件")

        diagnosis_heading.text = f"貳、{diagnosis_hospital_name}設備韌性備援診斷"
        self._ensure_new_page_before(introduction_heading)

        self._clear_section_between(introduction_heading, diagnosis_heading)
        self._insert_introduction_content(diagnosis_heading, hospital_name, report)
        self._ensure_new_page_before(diagnosis_heading)

        self._clear_section_between(diagnosis_heading, recommendation_heading)
        self._insert_diagnosis_content(
            document,
            recommendation_heading,
            hospital_name,
            report,
            normalized_document,
        )
        self._ensure_new_page_before(recommendation_heading)

        self._clear_section_between(recommendation_heading, attachment_heading)
        self._insert_recommendation_content(
            document=document,
            anchor_paragraph=attachment_heading,
            hospital_name=hospital_name,
            report=report,
            normalized_document=normalized_document,
        )
        self._move_section_break_paragraph_before_anchor(recommendation_heading, attachment_heading)
        self._ensure_new_page_before(attachment_heading)
        self._remove_leading_section_breaks_after(attachment_heading)

        self._render_attachment_section(document, attachment_heading, normalized_document.sources)
        self._apply_template_layout(document)

    def _render_generic_document(
        self,
        document: Document,
        project_name: str,
        report: ReportJSON,
        template_name: str,
        normalized_document: NormalizedDocument,
    ) -> None:
        self._add_title(document, title=report.title or project_name, template_name=template_name)
        self._add_heading(document, "摘要")
        self._add_paragraphs(document, [report.summary or "資料不足"])
        self._add_heading(document, "背景")
        self._add_paragraphs(document, [report.background or "資料不足"])
        self._render_report_tables(document, report.tables)
        self._add_heading(document, "建議事項")
        self._add_bullets(
            document,
            self._evidenced_text_list_to_strings(
                report.recommendations,
                require_evidence=True,
            ),
        )
        self._add_heading(document, "缺漏資料")
        self._add_bullets(document, report.missing_information)
        attachment_heading = self._add_heading(document, "附件")
        self._render_attachment_section(document, attachment_heading, normalized_document.sources)

    def _populate_cover_page(self, document: Document, hospital_name: str, report_date: str) -> None:
        cover_title = self._find_paragraph(
            document,
            lambda text: "重要急救責任醫院可行之維生系統設備韌性備援能力診斷輔導方案報告" in text,
        )
        cover_title.text = (
            "重要急救責任醫院可行之維生系統設備韌性備援能力診斷輔導方案報告"
            f"(醫院名稱：{hospital_name})"
        )

        cover_date = self._find_paragraph(document, lambda text: text.startswith("中華民國 "))
        cover_date.text = report_date

    def _populate_table_of_contents(self, document: Document, hospital_name: str) -> None:
        toc_entries_with_pages = [
            ("壹、前言", "2"),
            (f"貳、{hospital_name}設備韌性備援診斷", "3"),
            ("參、孤島效應下備援能力及建議", "4"),
            ("肆、附件", "6"),
        ]
        toc_entries = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "toc 1"][:4]
        for paragraph, (title, page_number) in zip(
            toc_entries,
            toc_entries_with_pages,
            strict=False,
        ):
            paragraph.text = f"{title}\t{page_number}"

    def _insert_introduction_content(
        self,
        anchor_paragraph: Paragraph,
        hospital_name: str,
        report: ReportJSON,
    ) -> None:
        paragraphs = self._build_fixed_introduction_paragraphs(hospital_name)
        if not paragraphs:
            self._insert_paragraph_before(
                anchor_paragraph,
                "資料不足",
                style_name=self.LIST_STYLE,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
                left_indent=self.INTRODUCTION_LEFT_INDENT,
                first_line_indent=self.INTRODUCTION_FIRST_LINE_INDENT,
            )
            return

        for index, paragraph_text in enumerate(paragraphs):
            self._insert_paragraph_before(
                anchor_paragraph,
                paragraph_text.strip(),
                style_name=self.LIST_STYLE,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
                left_indent=self.INTRODUCTION_LEFT_INDENT,
                first_line_indent=self.INTRODUCTION_FIRST_LINE_INDENT,
            )
            if index < len(paragraphs) - 1:
                self._insert_paragraph_before(
                    anchor_paragraph,
                    "",
                    style_name=self.LIST_STYLE,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
                    left_indent=self.INTRODUCTION_LEFT_INDENT,
                    first_line_indent=self.INTRODUCTION_FIRST_LINE_INDENT,
                )

    def _insert_diagnosis_content(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        hospital_name: str,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
        ) -> None:
        chapter_2_context = self._build_chapter_2_context(hospital_name, normalized_document)
        diagnosis_lead = self._build_diagnosis_lead_paragraph(chapter_2_context)
        self._insert_paragraph_before(
            anchor_paragraph,
            diagnosis_lead,
            style_name=self.LIST_STYLE,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            left_indent=self.DIAGNOSIS_BODY_LEFT_INDENT,
            first_line_indent=self.DIAGNOSIS_BODY_FIRST_LINE_INDENT,
            space_before=Pt(0),
            space_after=self.DIAGNOSIS_BODY_SPACE_AFTER,
        )

        diagnosis_image = self._collect_diagnosis_image_source(normalized_document)
        if diagnosis_image is not None:
            self._insert_image_before(
                anchor_paragraph,
                diagnosis_image,
                width=self.DIAGNOSIS_IMAGE_WIDTH,
            )
            self._insert_paragraph_before(
                anchor_paragraph,
                self._build_diagnosis_image_caption(chapter_2_context),
                style_name=self.LIST_STYLE,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            )

        paragraphs = self._build_diagnosis_paragraphs(
            report=report,
            chapter_2_context=chapter_2_context,
        )
        self._insert_paragraphs_before(
            anchor_paragraph,
            paragraphs,
            style_name=self.LIST_STYLE,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            left_indent=self.DIAGNOSIS_BODY_LEFT_INDENT,
            first_line_indent=self.DIAGNOSIS_BODY_FIRST_LINE_INDENT,
            space_before=Pt(0),
            space_after=self.DIAGNOSIS_BODY_SPACE_AFTER,
        )

        benefit_summary_rows = self._build_diagnosis_benefit_summary_rows(normalized_document)
        if benefit_summary_rows:
            self._insert_diagnosis_benefit_summary_table_before(
                document=document,
                anchor_paragraph=anchor_paragraph,
                summary_rows=benefit_summary_rows,
            )
            return

        diagnosis_tables = self._build_diagnosis_tables(normalized_document)
        if diagnosis_tables:
            self._insert_report_tables_before(
                document,
                anchor_paragraph,
                diagnosis_tables,
                include_titles=False,
            )

    def _build_diagnosis_paragraphs(
        self,
        report: ReportJSON,
        chapter_2_context: dict[str, str | bool],
    ) -> list[str]:
        ai_paragraphs: list[str] = []
        seen: set[str] = set()
        for paragraph in report.diagnosis_paragraphs:
            if not self._has_evidence(paragraph):
                continue
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            if text.startswith("根據訪視診斷") and "如問卷填答內容" in text:
                continue
            if text in seen:
                continue
            seen.add(text)
            ai_paragraphs.append(text)

        if ai_paragraphs:
            return ai_paragraphs[:2]

        return [
            "依現有資料尚無法完整判定該院設備韌性備援診斷總體觀察，需補充完整問卷、專家診斷表或補助效益調查資料後再行研判。"
        ]

    def _insert_recommendation_content(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        hospital_name: str,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> None:
        diagnosis_hospital_name = self._with_ministry_prefix(hospital_name)
        fixed_introduction = (
            f"根據前述，經本團隊診斷建議{diagnosis_hospital_name}在供電(供油量)、供水、供氣(氧氣)及資訊系統等"
            "項目之運作可強化有以下重點，以提升面對重大災難或孤島效性之運作量能。"
        )
        questionnaire_sections = self._build_questionnaire_recommendation_sections(normalized_document)
        introduction = (
            fixed_introduction
            if questionnaire_sections
            else report.recommendation_intro.strip() or fixed_introduction
        )
        self._insert_paragraph_before(
            anchor_paragraph,
            introduction,
            style_name=self.LIST_STYLE,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            left_indent=self.RECOMMENDATION_INTRO_LEFT_INDENT,
            first_line_indent=self.RECOMMENDATION_INTRO_FIRST_LINE_INDENT,
            space_before=Pt(0),
            space_after=self.RECOMMENDATION_INTRO_SPACE_AFTER,
        )

        self._insert_paragraph_before(
            anchor_paragraph,
            "",
            style_name=self.LIST_STYLE,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            left_indent=self.RECOMMENDATION_INTRO_LEFT_INDENT,
            first_line_indent=self.RECOMMENDATION_INTRO_FIRST_LINE_INDENT,
            space_before=Pt(0),
            space_after=self.RECOMMENDATION_INTRO_SPACE_AFTER,
        )

        sections = (
            questionnaire_sections
            if questionnaire_sections
            else self._build_complete_recommendation_sections(report, normalized_document)
        )
        for section_index, section in enumerate(sections, start=1):
            self._insert_paragraph_before(
                anchor_paragraph,
                self._format_numbered_section_title(section_index, section.title),
                style_name=self.LIST_STYLE,
                line_spacing=1.0,
                left_indent=self.RECOMMENDATION_SECTION_LEFT_INDENT,
                first_line_indent=self.RECOMMENDATION_SECTION_FIRST_LINE_INDENT,
            )
            use_numbering = len(section.paragraphs) > 1
            for paragraph_index, paragraph in enumerate(section.paragraphs, start=1):
                paragraph_text = self._evidenced_text_to_string(
                    paragraph,
                    require_evidence=False,
                )
                if not paragraph_text:
                    continue
                formatted_paragraph = (
                    self._format_numbered_subitem(paragraph_index, paragraph_text)
                    if use_numbering
                    else paragraph_text
                )
                layout_kwargs = self._get_recommendation_item_layout(paragraph_text)
                self._insert_paragraph_before(
                    anchor_paragraph,
                    formatted_paragraph,
                    style_name=self.LIST_STYLE,
                    line_spacing=1.0,
                    **layout_kwargs,
                )

        generated_chart_path = self.chart_generator.generate_pie_chart(normalized_document)
        analysis_images = self._collect_analysis_chart_sources(normalized_document)
        analysis_tables = self._build_analysis_tables(report, normalized_document)
        if generated_chart_path or analysis_images or analysis_tables:
            analysis_index = len(sections) + 1
            self._insert_paragraph_before(
                anchor_paragraph,
                self._format_numbered_section_title(analysis_index, "分析圖表"),
                style_name=self.LIST_STYLE,
                line_spacing=1.0,
                left_indent=self.RECOMMENDATION_SECTION_LEFT_INDENT,
                first_line_indent=self.RECOMMENDATION_SECTION_FIRST_LINE_INDENT,
            )
            if generated_chart_path is not None:
                self._insert_image_path_before(
                    anchor_paragraph,
                    generated_chart_path,
                    width=self.ANALYSIS_CHART_WIDTH,
                )
            elif analysis_images:
                for source in analysis_images:
                    self._insert_image_before(
                        anchor_paragraph,
                        source,
                        width=self.ANALYSIS_CHART_WIDTH,
                    )
            else:
                self._insert_report_tables_before(document, anchor_paragraph, analysis_tables)

        if report.missing_information:
            self._insert_paragraph_before(
                anchor_paragraph,
                "資料不足或待補充項目",
                style_name=self.LIST_STYLE,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            )
            self._insert_bullets_before(
                anchor_paragraph,
                report.missing_information,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
            )

    def _render_attachment_section(
        self,
        document: Document,
        attachment_heading: Paragraph,
        sources: list[ExtractedSource],
    ) -> None:
        self._remove_elements_after(attachment_heading._p)
        self._remove_leading_section_breaks_after(attachment_heading)

        if not sources:
            anchor_paragraph = self._find_insertion_anchor_after(document, attachment_heading)
            self._insert_paragraph_before(anchor_paragraph, "未提供附件。", style_name=self.LIST_STYLE)
            return

        composer = Composer(document)
        for index, source in enumerate(sources):
            if index > 0:
                composer.append(self._build_page_break_document())
            composer.append(self._build_attachment_document(source))

    def _apply_template_layout(self, document: Document) -> None:
        sections = list(document.sections)
        if not sections:
            return

        self._set_section_layout(
            sections[0],
            orientation=WD_ORIENT.PORTRAIT,
            page_width=Cm(21.0),
            page_height=Cm(29.7),
            top_margin=Cm(2.54),
            bottom_margin=Cm(2.54),
            left_margin=Cm(1.91),
            right_margin=Cm(1.25),
            header_distance=Cm(1.50),
            footer_distance=Cm(0.50),
        )

        if len(sections) >= 2:
            self._set_section_layout(
                sections[1],
                orientation=WD_ORIENT.PORTRAIT,
                page_width=Cm(21.0),
                page_height=Cm(29.7),
                top_margin=Cm(1.00),
                bottom_margin=Cm(2.54),
                left_margin=Cm(3.17),
                right_margin=Cm(3.17),
                header_distance=Cm(1.50),
                footer_distance=Cm(1.75),
            )

        for section in sections[2:]:
            self._set_section_layout(
                section,
                orientation=WD_ORIENT.LANDSCAPE,
                page_width=Cm(29.7),
                page_height=Cm(21.0),
                top_margin=Cm(0.83),
                bottom_margin=Cm(1.25),
                left_margin=Cm(1.50),
                right_margin=Cm(1.70),
                header_distance=Cm(1.50),
                footer_distance=Cm(1.75),
            )

    def _set_section_layout(
        self,
        section: object,
        *,
        orientation: WD_ORIENT,
        page_width: Length,
        page_height: Length,
        top_margin: Length,
        bottom_margin: Length,
        left_margin: Length,
        right_margin: Length,
        header_distance: Length,
        footer_distance: Length,
    ) -> None:
        section.orientation = orientation
        section.page_width = page_width
        section.page_height = page_height
        section.top_margin = top_margin
        section.bottom_margin = bottom_margin
        section.left_margin = left_margin
        section.right_margin = right_margin
        section.header_distance = header_distance
        section.footer_distance = footer_distance

    def _insert_source_content_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        source: ExtractedSource,
    ) -> None:
        if source.type == "docx":
            source_document = self._load_source_docx(source)
            if source_document is not None and self._copy_document_blocks_before(anchor_paragraph, source_document):
                return

        if source.type == "table":
            self._append_table_source_before(document, anchor_paragraph, source)
            return

        if source.type == "pdf":
            self._append_pdf_source_before(document, anchor_paragraph, source)
            return

        if source.type == "image":
            self._append_image_source_before(anchor_paragraph, source)
            return

        self._append_fallback_source_before(document, anchor_paragraph, source)

    def _copy_document_blocks_before(self, anchor_paragraph: Paragraph, source_document: Document) -> bool:
        inserted_any = False
        body = source_document.element.body
        for child in body.iterchildren():
            local_name = child.tag.split("}")[-1]
            if local_name not in {"p", "tbl"}:
                continue
            if self._element_contains_section_break(child):
                continue
            anchor_paragraph._p.addprevious(deepcopy(child))
            inserted_any = True
        return inserted_any

    def _append_table_source_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        source: ExtractedSource,
    ) -> None:
        for table_index, extracted_table in enumerate(source.tables, start=1):
            title = self._normalize_attachment_table_title(
                title=extracted_table.sheet_name or extracted_table.title,
                source_name=source.source_name,
                fallback=f"表格 {table_index}",
            )
            self._insert_paragraph_before(anchor_paragraph, title, style_name="Normal", bold=True)
            if extracted_table.raw_rows:
                self._insert_raw_table_before(document, anchor_paragraph, extracted_table.raw_rows)
            else:
                self._insert_extracted_table_before(document, anchor_paragraph, extracted_table)

    def _append_pdf_source_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        source: ExtractedSource,
    ) -> None:
        if source.text:
            for block in self._split_blocks(source.text, limit=120):
                self._insert_paragraph_before(anchor_paragraph, block, style_name="Normal")

        for table_index, extracted_table in enumerate(source.tables, start=1):
            title = self._normalize_attachment_table_title(
                title=extracted_table.title,
                source_name=source.source_name,
                fallback=f"PDF 表格 {table_index}",
            )
            self._insert_paragraph_before(anchor_paragraph, title, style_name="Normal", bold=True)
            self._insert_extracted_table_before(document, anchor_paragraph, extracted_table)

    def _append_image_source_before(self, anchor_paragraph: Paragraph, source: ExtractedSource) -> None:
        source_path_value = source.metadata.get("path")
        if source_path_value:
            source_path = Path(source_path_value)
            if source_path.is_file():
                paragraph = anchor_paragraph.insert_paragraph_before("")
                run = paragraph.add_run()
                run.add_picture(str(source_path), width=Inches(6.0))
                return

        self._insert_paragraph_before(
            anchor_paragraph,
            "此附件為圖片檔，但目前無法直接嵌入原圖。",
            style_name="Normal",
        )

    def _append_fallback_source_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        source: ExtractedSource,
    ) -> None:
        if source.paragraphs:
            for paragraph in source.paragraphs:
                self._insert_paragraph_before(anchor_paragraph, paragraph, style_name="Normal")
        elif source.text:
            for block in self._split_blocks(source.text, limit=60):
                self._insert_paragraph_before(anchor_paragraph, block, style_name="Normal")
        else:
            self._insert_paragraph_before(
                anchor_paragraph,
                "附件內容無可直接貼上的文字。",
                style_name="Normal",
            )

        for table_index, extracted_table in enumerate(source.tables, start=1):
            title = self._normalize_attachment_table_title(
                title=extracted_table.title,
                source_name=source.source_name,
                fallback=f"表格 {table_index}",
            )
            self._insert_paragraph_before(anchor_paragraph, title, style_name="Normal", bold=True)
            self._insert_extracted_table_before(document, anchor_paragraph, extracted_table)

    def _build_page_break_document(self) -> Document:
        page_break_document = Document()
        page_break_document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        return page_break_document

    def _build_attachment_document(self, source: ExtractedSource) -> Document:
        if source.type == "docx":
            source_document = self._load_source_docx(source)
            if source_document is not None:
                return source_document

        attachment_document = Document()

        if source.type == "table":
            self._append_table_source_document(attachment_document, source)
        elif source.type == "pdf":
            self._append_pdf_source_document(attachment_document, source)
        elif source.type == "image":
            self._append_image_source_document(attachment_document, source)
        else:
            self._append_fallback_source_document(attachment_document, source)

        self._trim_empty_attachment_paragraphs(attachment_document)
        return attachment_document

    def _append_table_source_document(
        self,
        document: Document,
        source: ExtractedSource,
    ) -> None:
        if not source.tables:
            self._append_paragraph(document, "資料不足", style_name="Normal")
            return

        show_titles = len(source.tables) > 1
        for table_index, extracted_table in enumerate(source.tables, start=1):
            if show_titles:
                title = self._normalize_attachment_table_title(
                    title=extracted_table.sheet_name or extracted_table.title,
                    source_name=source.source_name,
                    fallback=f"表格 {table_index}",
                )
                self._append_paragraph(document, title, style_name="Normal", bold=True)

            if extracted_table.raw_rows:
                self._append_raw_table_document(document, extracted_table.raw_rows)
            else:
                self._append_extracted_table_document(document, extracted_table)

    def _append_pdf_source_document(
        self,
        document: Document,
        source: ExtractedSource,
    ) -> None:
        if source.text:
            for block in self._split_blocks(source.text, limit=120):
                self._append_paragraph(document, block, style_name="Normal")

        show_titles = len(source.tables) > 1
        for table_index, extracted_table in enumerate(source.tables, start=1):
            if show_titles:
                title = self._normalize_attachment_table_title(
                    title=extracted_table.title,
                    source_name=source.source_name,
                    fallback=f"PDF 表格 {table_index}",
                )
                self._append_paragraph(document, title, style_name="Normal", bold=True)
            self._append_extracted_table_document(document, extracted_table)

    def _append_image_source_document(self, document: Document, source: ExtractedSource) -> None:
        source_path_value = source.metadata.get("path")
        if source_path_value:
            source_path = Path(source_path_value)
            if source_path.is_file():
                paragraph = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
                paragraph.add_run().add_picture(str(source_path), width=Inches(6.0))
                return

        self._append_paragraph(
            document,
            "附件圖片無法直接貼上原始檔案，請檢查來源路徑。",
            style_name="Normal",
        )

    def _append_fallback_source_document(
        self,
        document: Document,
        source: ExtractedSource,
    ) -> None:
        if source.paragraphs:
            for paragraph in source.paragraphs:
                self._append_paragraph(document, paragraph, style_name="Normal")
        elif source.text:
            for block in self._split_blocks(source.text, limit=60):
                self._append_paragraph(document, block, style_name="Normal")
        else:
            self._append_paragraph(document, "附件內容無可直接貼上的文字。", style_name="Normal")

        show_titles = len(source.tables) > 1
        for table_index, extracted_table in enumerate(source.tables, start=1):
            if show_titles:
                title = self._normalize_attachment_table_title(
                    title=extracted_table.title,
                    source_name=source.source_name,
                    fallback=f"表格 {table_index}",
                )
                self._append_paragraph(document, title, style_name="Normal", bold=True)
            self._append_extracted_table_document(document, extracted_table)

    def _append_extracted_table_document(
        self,
        document: Document,
        extracted_table: ExtractedTable,
    ) -> None:
        headers, rows = self._normalize_extracted_table(extracted_table)
        if not headers and not rows:
            self._append_paragraph(document, "資料不足", style_name="Normal")
            return

        column_count = len(headers) if headers else max(len(row) for row in rows)
        table = document.add_table(rows=1, cols=max(column_count, 1))
        table.style = "Table Grid"

        if headers:
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
        else:
            table.rows[0].cells[0].text = "資料"

        for row in rows:
            row_cells = table.add_row().cells
            normalized_row = row + [""] * (len(row_cells) - len(row))
            for index, value in enumerate(normalized_row[: len(row_cells)]):
                row_cells[index].text = value

    def _append_raw_table_document(
        self,
        document: Document,
        raw_rows: list[list[object]],
    ) -> None:
        if not raw_rows:
            self._append_paragraph(document, "資料不足", style_name="Normal")
            return

        column_count = max(len(row) for row in raw_rows)
        table = document.add_table(rows=1, cols=max(column_count, 1))
        table.style = "Table Grid"

        for row_index, row in enumerate(raw_rows):
            row_cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
            normalized_row = row + [None] * (len(row_cells) - len(row))
            for index, value in enumerate(normalized_row[: len(row_cells)]):
                row_cells[index].text = self._stringify_value(value)

    def _trim_empty_attachment_paragraphs(self, document: Document) -> None:
        paragraphs = list(document.paragraphs)
        for paragraph in paragraphs:
            has_text = bool(self._normalize_text(paragraph.text))
            has_drawing = paragraph._p.find(".//w:drawing", self.WORD_NS) is not None
            has_page_break = paragraph._p.find(".//w:br", self.WORD_NS) is not None
            if has_text or has_drawing or has_page_break:
                break
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)

    def _clear_section_between(self, start_heading: Paragraph, end_heading: Paragraph) -> None:
        self._remove_elements_between(start_heading._p, end_heading._p)

    def _remove_elements_between(self, start_element: object, end_element: object) -> None:
        current = start_element.getnext()
        while current is not None and current is not end_element:
            next_element = current.getnext()
            if self._element_contains_section_break(current):
                current = next_element
                continue
            current.getparent().remove(current)
            current = next_element

    def _remove_elements_after(self, start_element: object) -> None:
        current = start_element.getnext()
        while current is not None:
            next_element = current.getnext()
            if current.tag == qn("w:sectPr") or self._element_contains_section_break(current):
                current = next_element
                continue
            current.getparent().remove(current)
            current = next_element

    def _element_contains_section_break(self, element: object) -> bool:
        return element.find(".//w:sectPr", self.WORD_NS) is not None

    def _find_paragraph(self, document: Document, predicate: object) -> Paragraph:
        for paragraph in document.paragraphs:
            normalized = self._normalize_text(paragraph.text)
            if normalized and predicate(normalized):
                return paragraph
        raise ValueError("找不到對應的模板段落。")

    def _find_heading_paragraph(self, document: Document, predicate: object) -> Paragraph:
        for paragraph in document.paragraphs:
            if not paragraph.style.name.startswith("Heading"):
                continue
            normalized = self._normalize_text(paragraph.text)
            if normalized and predicate(normalized):
                return paragraph
        raise ValueError("找不到對應的模板章節標題。")

    def _find_insertion_anchor_after(self, document: Document, paragraph: Paragraph) -> Paragraph:
        current = paragraph._p.getnext()
        while current is not None:
            if current.tag == qn("w:p"):
                return Paragraph(current, paragraph._parent)
            current = current.getnext()
        return document.add_paragraph()

    def _insert_page_break_before(self, anchor_paragraph: Paragraph) -> None:
        paragraph = anchor_paragraph.insert_paragraph_before("")
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def _ensure_new_page_before(self, paragraph: Paragraph) -> None:
        previous = paragraph._p.getprevious()
        while previous is not None and previous.tag != qn("w:p"):
            if previous.tag == qn("w:sectPr"):
                return
            previous = previous.getprevious()

        if previous is None:
            return

        if self._element_contains_section_break(previous) or self._element_contains_page_break(previous):
            return

        self._insert_page_break_before(paragraph)

    def _element_contains_page_break(self, element: object) -> bool:
        return element.find('.//w:br[@w:type="page"]', self.WORD_NS) is not None

    def _move_section_break_paragraph_before_anchor(
        self,
        start_heading: Paragraph,
        anchor_paragraph: Paragraph,
    ) -> None:
        current = start_heading._p.getnext()
        while current is not None and current is not anchor_paragraph._p:
            next_element = current.getnext()
            if current.tag == qn("w:p") and self._element_contains_section_break(current):
                current.getparent().remove(current)
                anchor_paragraph._p.addprevious(current)
                return
            current = next_element

    def _remove_leading_section_breaks_after(self, paragraph: Paragraph) -> None:
        current = paragraph._p.getnext()
        while current is not None:
            next_element = current.getnext()
            if current.tag != qn("w:p") or not self._element_contains_section_break(current):
                return
            current.getparent().remove(current)
            current = next_element

    def _insert_image_before(
        self,
        anchor_paragraph: Paragraph,
        source: ExtractedSource,
        width: Length,
    ) -> None:
        source_path_value = source.metadata.get("path")
        if not source_path_value:
            return

        source_path = Path(source_path_value)
        if not source_path.is_file():
            return

        paragraph = anchor_paragraph.insert_paragraph_before("")
        paragraph.add_run().add_picture(str(source_path), width=width)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_image_path_before(
        self,
        anchor_paragraph: Paragraph,
        image_path: Path,
        width: Length,
    ) -> None:
        if not image_path.is_file():
            return

        paragraph = anchor_paragraph.insert_paragraph_before("")
        paragraph.add_run().add_picture(str(image_path), width=width)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_paragraph_before(
        self,
        anchor_paragraph: Paragraph,
        text: str,
        style_name: str | None = None,
        bold: bool = False,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        line_spacing: float | Pt | None = None,
        left_indent: Length | None = None,
        right_indent: Length | None = None,
        first_line_indent: Length | None = None,
        space_before: Length | None = None,
        space_after: Length | None = None,
    ) -> Paragraph:
        paragraph = anchor_paragraph.insert_paragraph_before("")
        if style_name:
            paragraph.style = style_name
        run = paragraph.add_run(text)
        run.bold = bold
        self._apply_run_font(run)
        self._apply_paragraph_layout(
            paragraph,
            alignment=alignment,
            line_spacing=line_spacing,
            left_indent=left_indent,
            right_indent=right_indent,
            first_line_indent=first_line_indent,
            space_before=space_before,
            space_after=space_after,
        )
        return paragraph

    def _insert_paragraphs_before(
        self,
        anchor_paragraph: Paragraph,
        paragraphs: list[str],
        style_name: str,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        line_spacing: float | Pt | None = None,
        left_indent: Length | None = None,
        right_indent: Length | None = None,
        first_line_indent: Length | None = None,
        space_before: Length | None = None,
        space_after: Length | None = None,
    ) -> None:
        if not paragraphs:
            self._insert_paragraph_before(
                anchor_paragraph,
                "資料不足",
                style_name=style_name,
                alignment=alignment,
                line_spacing=line_spacing,
                left_indent=left_indent,
                right_indent=right_indent,
                first_line_indent=first_line_indent,
                space_before=space_before,
                space_after=space_after,
            )
            return

        for paragraph in paragraphs:
            text = paragraph.strip()
            if not text:
                continue
            self._insert_paragraph_before(
                anchor_paragraph,
                text,
                style_name=style_name,
                alignment=alignment,
                line_spacing=line_spacing,
                left_indent=left_indent,
                right_indent=right_indent,
                first_line_indent=first_line_indent,
                space_before=space_before,
                space_after=space_after,
            )

    def _insert_bullets_before(
        self,
        anchor_paragraph: Paragraph,
        items: list[str],
        line_spacing: float | Pt | None = None,
    ) -> None:
        if not items:
            self._insert_paragraph_before(
                anchor_paragraph,
                "資料不足",
                style_name=self.BULLET_STYLE,
                line_spacing=line_spacing,
            )
            return

        for item in items:
            self._insert_paragraph_before(
                anchor_paragraph,
                item,
                style_name=self.BULLET_STYLE,
                line_spacing=line_spacing,
            )

    def _apply_paragraph_layout(
        self,
        paragraph: Paragraph,
        *,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        line_spacing: float | Pt | None = None,
        left_indent: Length | None = None,
        right_indent: Length | None = None,
        first_line_indent: Length | None = None,
        space_before: Length | None = None,
        space_after: Length | None = None,
    ) -> None:
        if alignment is not None:
            paragraph.alignment = alignment
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing
        if left_indent is not None:
            paragraph.paragraph_format.left_indent = left_indent
        if right_indent is not None:
            paragraph.paragraph_format.right_indent = right_indent
        if first_line_indent is not None:
            paragraph.paragraph_format.first_line_indent = first_line_indent
        if space_before is not None:
            paragraph.paragraph_format.space_before = space_before
        if space_after is not None:
            paragraph.paragraph_format.space_after = space_after

    def _apply_run_font(
        self,
        run: object,
        *,
        font_name: str | None = None,
        font_size: Pt | None = None,
        bold: bool | None = None,
    ) -> None:
        resolved_font_name = font_name or self.BODY_FONT_NAME
        resolved_font_size = font_size or self.BODY_FONT_SIZE
        run.font.name = resolved_font_name
        run.font.size = resolved_font_size
        if bold is not None:
            run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for font_slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{font_slot}"), resolved_font_name)

    def _apply_document_font(self, document: Document) -> None:
        self._apply_font_to_paragraphs(document.paragraphs)
        self._apply_font_to_tables(document.tables)
        for section in document.sections:
            self._apply_font_to_paragraphs(section.header.paragraphs)
            self._apply_font_to_tables(section.header.tables)
            self._apply_font_to_paragraphs(section.footer.paragraphs)
            self._apply_font_to_tables(section.footer.tables)

    def _apply_table_of_contents_format(self, document: Document) -> None:
        toc_title = self._find_toc_title_paragraph(document)
        if toc_title is not None:
            self._apply_paragraph_runs_font(
                toc_title,
                font_name=self.BODY_FONT_NAME,
                font_size=self.TOC_TITLE_FONT_SIZE,
                bold=True,
            )

        toc_entries = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "toc 1"][:4]
        for paragraph in toc_entries:
            self._apply_paragraph_runs_font(
                paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=self.TOC_ENTRY_FONT_SIZE,
                bold=True,
            )

    def _apply_cover_page_format(self, document: Document) -> None:
        ministry_paragraph = self._find_first_paragraph_by_text(document, "衛生福利部")
        if ministry_paragraph is not None:
            self._apply_paragraph_runs_font(
                ministry_paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=self.COVER_MINISTRY_FONT_SIZE,
                bold=False,
            )

        plan_paragraph = self._find_first_paragraph_by_text(
            document,
            "114-115年度「急救責任醫院維生系統韌性盤點暨輔導計畫」",
        )
        if plan_paragraph is not None:
            self._apply_paragraph_runs_font(
                plan_paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=self.COVER_PLAN_FONT_SIZE,
                bold=False,
            )

        report_title_paragraph = self._find_first_paragraph_by_prefix(
            document,
            "重要急救責任醫院可行之維生系統設備韌性備援能力診斷輔導方案報告",
        )
        if report_title_paragraph is not None:
            self._apply_paragraph_runs_font(
                report_title_paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=self.COVER_REPORT_TITLE_FONT_SIZE,
                bold=False,
            )

    def _apply_recommendation_chapter_font(self, document: Document) -> None:
        recommendation_heading = self._find_heading_paragraph(
            document,
            lambda text: text == "參、孤島效應下備援能力及建議",
        )
        attachment_heading = self._find_heading_paragraph(document, lambda text: text == "肆、附件")
        if recommendation_heading is None or attachment_heading is None:
            return

        paragraphs = document.paragraphs
        start_index = self._get_paragraph_index_by_element(paragraphs, recommendation_heading)
        end_index = self._get_paragraph_index_by_element(paragraphs, attachment_heading)
        if start_index is None or end_index is None:
            return
        recommendation_started = False

        for paragraph in paragraphs[start_index + 1 : end_index]:
            text = self._normalize_text(paragraph.text)
            if not text:
                continue

            if self._is_numbered_recommendation_section(text):
                recommendation_started = True

            target_font_size = (
                self.RECOMMENDATION_BODY_FONT_SIZE
                if recommendation_started
                else self.RECOMMENDATION_INTRO_FONT_SIZE
            )
            self._apply_paragraph_font_size(
                paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=target_font_size,
            )
            if not recommendation_started:
                self._apply_paragraph_layout(
                    paragraph,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
                    left_indent=self.RECOMMENDATION_INTRO_LEFT_INDENT,
                    first_line_indent=self.RECOMMENDATION_INTRO_FIRST_LINE_INDENT,
                    space_before=Pt(0),
                    space_after=self.RECOMMENDATION_INTRO_SPACE_AFTER,
                )

    def _is_numbered_recommendation_section(self, text: str) -> bool:
        return bool(re.match(r"^[一二三四五六七八九十]+、", text))

    def _apply_forced_layout_overrides(self, document: Document) -> None:
        self._force_recommendation_intro_layout(document)

    def _force_recommendation_intro_layout(self, document: Document) -> None:
        recommendation_heading = self._find_heading_paragraph(
            document,
            lambda text: "孤島效應下備援能力及建議" in text,
        )
        attachment_heading = self._find_heading_paragraph(
            document,
            lambda text: text == "肆、附件" or text == "附件",
        )
        if recommendation_heading is None or attachment_heading is None:
            return

        paragraphs = document.paragraphs
        start_index = self._get_paragraph_index_by_element(paragraphs, recommendation_heading)
        end_index = self._get_paragraph_index_by_element(paragraphs, attachment_heading)
        if start_index is None or end_index is None:
            return

        for paragraph in paragraphs[start_index + 1 : end_index]:
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            if self._is_numbered_recommendation_section(text):
                return

            self._apply_paragraph_font_size(
                paragraph,
                font_name=self.BODY_FONT_NAME,
                font_size=self.RECOMMENDATION_INTRO_FONT_SIZE,
            )
            self._apply_paragraph_layout(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing=self.BODY_PARAGRAPH_LINE_SPACING,
                left_indent=self.RECOMMENDATION_INTRO_LEFT_INDENT,
                first_line_indent=self.RECOMMENDATION_INTRO_FIRST_LINE_INDENT,
                space_before=Pt(0),
                space_after=self.RECOMMENDATION_INTRO_SPACE_AFTER,
            )

    def _get_paragraph_index_by_element(
        self,
        paragraphs: list[Paragraph],
        target_paragraph: Paragraph,
    ) -> int | None:
        for index, paragraph in enumerate(paragraphs):
            if paragraph._p is target_paragraph._p:
                return index
        return None

    def _find_first_paragraph_by_text(self, document: Document, text: str) -> Paragraph | None:
        target = self._normalize_text(text)
        for paragraph in document.paragraphs:
            if self._normalize_text(paragraph.text) == target:
                return paragraph
        return None

    def _find_first_paragraph_by_prefix(self, document: Document, prefix: str) -> Paragraph | None:
        normalized_prefix = self._normalize_text(prefix)
        for paragraph in document.paragraphs:
            if self._normalize_text(paragraph.text).startswith(normalized_prefix):
                return paragraph
        return None

    def _find_toc_title_paragraph(self, document: Document) -> Paragraph | None:
        title = self._find_first_paragraph_by_text(document, "目錄")
        if title is not None:
            return title

        toc_entries = [paragraph for paragraph in document.paragraphs if paragraph.style.name == "toc 1"]
        if not toc_entries:
            return None

        first_entry = toc_entries[0]
        paragraphs = document.paragraphs
        entry_index = paragraphs.index(first_entry)
        for paragraph in reversed(paragraphs[:entry_index]):
            if self._normalize_text(paragraph.text):
                return paragraph
        return None

    def _apply_paragraph_runs_font(
        self,
        paragraph: Paragraph,
        *,
        font_name: str,
        font_size: Pt,
        bold: bool,
    ) -> None:
        for run in paragraph.runs:
            self._apply_run_font(
                run,
                font_name=font_name,
                font_size=font_size,
                bold=bold,
            )

    def _apply_paragraph_font_size(
        self,
        paragraph: Paragraph,
        *,
        font_name: str,
        font_size: Pt,
    ) -> None:
        for run in paragraph.runs:
            self._apply_run_font(
                run,
                font_name=font_name,
                font_size=font_size,
            )

    def _apply_font_to_paragraphs(self, paragraphs: list[Paragraph]) -> None:
        for paragraph in paragraphs:
            for run in paragraph.runs:
                self._apply_run_font(run)

    def _apply_font_to_tables(self, tables: list[Table]) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    self._apply_font_to_paragraphs(cell.paragraphs)
                    if cell.tables:
                        self._apply_font_to_tables(cell.tables)

    def _apply_table_font_size(self, table: Table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        self._apply_run_font(run)

    def _get_recommendation_item_layout(self, paragraph: str) -> dict[str, Length]:
        stripped = paragraph.strip()
        if re.match(r"^[A-ZＡ-Ｚ]:", stripped) or stripped.startswith("故A/[B*C]"):
            return {
                "left_indent": self.RECOMMENDATION_FORMULA_LEFT_INDENT,
                "first_line_indent": self.RECOMMENDATION_FORMULA_FIRST_LINE_INDENT,
            }

        return {
            "left_indent": self.RECOMMENDATION_ITEM_LEFT_INDENT,
            "first_line_indent": self.RECOMMENDATION_ITEM_FIRST_LINE_INDENT,
        }

    def _insert_report_tables_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        tables: list[ReportTable],
        include_titles: bool = True,
    ) -> None:
        for report_table in tables:
            if include_titles:
                title = self._normalize_analysis_table_title(report_table.title)
                self._insert_paragraph_before(anchor_paragraph, title, style_name="Normal", bold=True)
            self._insert_report_table_before(document, anchor_paragraph, report_table)

    def _insert_report_table_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        report_table: ReportTable,
    ) -> None:
        headers = report_table.columns or ["欄位"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for row in report_table.rows:
            row_cells = table.add_row().cells
            normalized_row = row + [""] * (len(headers) - len(row))
            for index, value in enumerate(normalized_row[: len(headers)]):
                row_cells[index].text = value

        self._apply_table_font_size(table)
        anchor_paragraph._p.addprevious(table._tbl)

    def _build_diagnosis_benefit_summary_rows(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[list[str]]:
        for source in normalized_document.sources:
            if source.type != "table" or not self._is_benefit_table_source(source):
                continue

            for table in source.tables:
                rows = [row for row in table.rows if isinstance(row, dict)]
                if not rows:
                    continue

                summary_rows: list[list[str]] = []
                for row in rows:
                    category = self._stringify_value(self._find_row_value(row, "類別")).strip()
                    if not category or category.startswith("以上"):
                        continue

                    before_value = self._format_benefit_summary_value(self._find_row_value(row, "補助前"))
                    after_value = self._format_benefit_summary_value(self._find_row_value(row, "補助後"))
                    delta_value = self._format_benefit_summary_value(self._find_row_value(row, "增減變化"))
                    attained_value = self._format_benefit_attainment_value(
                        self._find_row_value(row, "孤島效應達標")
                    )
                    focus_value = self._format_benefit_note_value(self._find_row_value(row, "補助重點"))
                    benefit_value = self._format_benefit_note_value(self._find_row_value(row, "補助效益"))
                    note_value = self._format_benefit_note_value(self._find_row_value(row, "備註"))

                    summary_rows.append(
                        [
                            category,
                            before_value,
                            after_value,
                            delta_value,
                            attained_value,
                            focus_value,
                            benefit_value,
                            note_value,
                        ]
                    )

                if summary_rows:
                    return summary_rows

        return []

    def _insert_diagnosis_benefit_summary_table_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        summary_rows: list[list[str]],
    ) -> None:
        table = document.add_table(rows=2 + len(summary_rows), cols=8)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        top_headers = ["", "天數", "天數", "增減變化天數", "孤島效應達標", "補助重點", "補助效益", "備註"]
        bottom_headers = ["類別", "補助前", "補助後", "", "", "", "", ""]

        for column_index, value in enumerate(top_headers):
            self._set_table_cell_text(
                table.rows[0].cells[column_index],
                value,
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for column_index, value in enumerate(bottom_headers):
            self._set_table_cell_text(
                table.rows[1].cells[column_index],
                value,
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for column_index in range(3, 8):
            merged_cell = table.rows[0].cells[column_index].merge(table.rows[1].cells[column_index])
            self._set_table_cell_text(
                merged_cell,
                top_headers[column_index],
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for row_index, row_values in enumerate(summary_rows, start=2):
            row_cells = table.rows[row_index].cells
            for column_index, value in enumerate(row_values):
                fill = self.DIAGNOSIS_SUMMARY_HIGHLIGHT_FILL if column_index == 1 else None
                alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index <= 4 else WD_ALIGN_PARAGRAPH.LEFT
                self._set_table_cell_text(
                    row_cells[column_index],
                    value,
                    font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                    bold=False,
                    font_color=RGBColor(0, 0, 0),
                    alignment=alignment,
                    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                    fill=fill,
                )

        anchor_paragraph._p.addprevious(table._tbl)

    def _set_table_cell_text(
        self,
        cell: object,
        text: str,
        *,
        font_size: Pt,
        bold: bool,
        font_color: RGBColor,
        alignment: WD_ALIGN_PARAGRAPH,
        vertical_alignment: WD_CELL_VERTICAL_ALIGNMENT,
        fill: str | None = None,
    ) -> None:
        cell.text = text
        cell.vertical_alignment = vertical_alignment
        if fill:
            self._set_cell_fill(cell, fill)

        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                self._apply_run_font(
                    run,
                    font_name=self.BODY_FONT_NAME,
                    font_size=font_size,
                    bold=bold,
                )
                run.font.color.rgb = font_color

    def _set_cell_fill(self, cell: object, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    def _format_benefit_summary_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""

        numeric_value = self._to_float(text)
        if numeric_value is None:
            return text

        if numeric_value.is_integer():
            return str(int(numeric_value))
        return f"{numeric_value:.2f}".rstrip("0").rstrip(".")

    def _format_benefit_attainment_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""

        normalized = text.lower()
        if normalized in {"v", "y", "yes", "true", "達標"}:
            return "達標"
        if normalized in {"x", "n", "no", "false", "未達標"}:
            return "未達標"
        return text

    def _format_benefit_note_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""
        return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()

    def _insert_extracted_table_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        extracted_table: ExtractedTable,
    ) -> None:
        headers, rows = self._normalize_extracted_table(extracted_table)
        if not headers and not rows:
            self._insert_paragraph_before(anchor_paragraph, "資料不足", style_name="Normal")
            return

        column_count = len(headers) if headers else max(len(row) for row in rows)
        table = document.add_table(rows=1, cols=max(column_count, 1))
        table.style = "Table Grid"

        if headers:
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
        else:
            table.rows[0].cells[0].text = "內容"

        for row in rows:
            row_cells = table.add_row().cells
            normalized_row = row + [""] * (len(row_cells) - len(row))
            for index, value in enumerate(normalized_row[: len(row_cells)]):
                row_cells[index].text = value

        self._apply_table_font_size(table)

        self._apply_table_font_size(table)
        anchor_paragraph._p.addprevious(table._tbl)

    def _insert_raw_table_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        raw_rows: list[list[object]],
    ) -> None:
        if not raw_rows:
            self._insert_paragraph_before(anchor_paragraph, "資料不足", style_name="Normal")
            return

        column_count = max(len(row) for row in raw_rows)
        table = document.add_table(rows=1, cols=max(column_count, 1))
        table.style = "Table Grid"

        for row_index, row in enumerate(raw_rows):
            row_cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
            normalized_row = row + [None] * (len(row_cells) - len(row))
            for index, value in enumerate(normalized_row[: len(row_cells)]):
                row_cells[index].text = self._stringify_value(value)

        self._apply_table_font_size(table)

        self._apply_table_font_size(table)
        anchor_paragraph._p.addprevious(table._tbl)

    def _build_diagnosis_paragraphs_legacy(self, report: ReportJSON) -> list[str]:
        paragraphs = self._clean_text_items(
            self._evidenced_text_list_to_strings(
                report.diagnosis_paragraphs,
                require_evidence=True,
            )
            or self._fallback_diagnosis_paragraphs(report)
        )

        if len(paragraphs) < 4 and report.summary.strip():
            paragraphs.append(report.summary.strip())

        if len(paragraphs) < 4 and report.key_findings:
            evidence_lines = [
                f"{finding.item}：{self._format_evidence_references(finding.evidence)}"
                for finding in report.key_findings[:2]
                if finding.evidence
            ]
            if evidence_lines:
                paragraphs.append("本次診斷可參照之主要依據包括：" + "；".join(evidence_lines) + "。")

        return self._clean_text_items(paragraphs)

    def _build_complete_recommendation_sections(
        self,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> list[ReportSection]:
        questionnaire_sections = self._build_questionnaire_recommendation_sections(normalized_document)
        if questionnaire_sections:
            return questionnaire_sections

        sections = report.recommendation_sections or self._build_recommendation_sections(report)
        backlog = self._evidenced_text_list_to_strings(
            report.recommendations,
            require_evidence=True,
        )
        grounded_sections = self._build_grounded_recommendation_sections(
            report=report,
            normalized_document=normalized_document,
        )
        grouped_sections: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }
        other_sections: list[ReportSection] = []

        for section in sections:
            normalized_title = self._normalize_recommendation_section_title(section.title)
            paragraphs = self._clean_text_items(
                self._evidenced_text_list_to_strings(
                    section.paragraphs,
                    require_evidence=True,
                )
            )
            while len(paragraphs) < 2 and backlog:
                candidate = backlog.pop(0)
                if candidate not in paragraphs:
                    paragraphs.append(candidate)
            if normalized_title in grouped_sections:
                for paragraph in paragraphs:
                    if paragraph not in grouped_sections[normalized_title]:
                        grouped_sections[normalized_title].append(paragraph)
            elif paragraphs:
                other_sections.append(
                    ReportSection(title=section.title.strip(), paragraphs=paragraphs)
                )

        for title, grounded_paragraphs in grounded_sections.items():
            if len(grouped_sections[title]) >= 2:
                continue
            for paragraph in grounded_paragraphs:
                if paragraph not in grouped_sections[title]:
                    grouped_sections[title].append(paragraph)
                if len(grouped_sections[title]) >= 2:
                    break

        if backlog:
            for candidate in backlog:
                if candidate not in grouped_sections["其他(工安衛)"]:
                    grouped_sections["其他(工安衛)"].append(candidate)

        completed = [
            ReportSection(title=title, paragraphs=self._clean_text_items(paragraphs))
            for title, paragraphs in grouped_sections.items()
            if paragraphs
        ]
        completed.extend(other_sections)
        return completed

    def _build_questionnaire_recommendation_sections(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[ReportSection]:
        grouped: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }

        for source in normalized_document.sources:
            if not self._is_questionnaire_recommendation_source(source):
                continue

            for table in source.tables:
                if not table.semantic_lines:
                    continue

                power_items = self._extract_questionnaire_recommendation_items(
                    table.semantic_lines,
                    "二、供電相關",
                )
                water_items = self._extract_questionnaire_recommendation_items(
                    table.semantic_lines,
                    "三、供水相關",
                )
                gas_items = self._extract_questionnaire_recommendation_items(
                    table.semantic_lines,
                    "四、醫用氣體",
                )
                it_items = self._extract_questionnaire_recommendation_items(
                    table.semantic_lines,
                    "五、資訊系統備援",
                )
                monitoring_items = self._extract_questionnaire_recommendation_items(
                    table.semantic_lines,
                    "七、中央監控系統",
                )

                for item in power_items:
                    if item not in grouped["供電部分"]:
                        grouped["供電部分"].append(item)
                for item in water_items:
                    if item not in grouped["供水部分"]:
                        grouped["供水部分"].append(item)
                for item in [*gas_items, *it_items]:
                    if item not in grouped["供氧氣及IT系統部分之建議"]:
                        grouped["供氧氣及IT系統部分之建議"].append(item)
                for item in monitoring_items:
                    if item not in grouped["其他(工安衛)"]:
                        grouped["其他(工安衛)"].append(item)

        return [
            ReportSection(title=title, paragraphs=paragraphs)
            for title, paragraphs in grouped.items()
            if paragraphs
        ]

    def _is_questionnaire_recommendation_source(self, source: ExtractedSource) -> bool:
        if source.type != "docx":
            return False
        return self._has_questionnaire_recommendation_content(source)

    def _has_questionnaire_recommendation_content(self, source: ExtractedSource) -> bool:
        target_sections = (
            "診斷表 | 二、供電相關 | 建議",
            "診斷表 | 三、供水相關 | 建議",
            "診斷表 | 四、醫用氣體 | 建議",
            "診斷表 | 五、資訊系統備援 | 建議",
            "診斷表 | 七、中央監控系統 | 建議",
        )
        for table in source.tables:
            for line in table.semantic_lines:
                normalized_line = self._normalize_text(line)
                if any(normalized_line.startswith(prefix) for prefix in target_sections):
                    return True
        return False

    def _extract_questionnaire_recommendation_items(
        self,
        semantic_lines: list[str],
        section_label: str,
    ) -> list[str]:
        detailed_items: list[str] = []
        general_item = ""

        for line in semantic_lines:
            normalized_line = self._normalize_text(line)
            if not normalized_line.startswith(f"診斷表 | {section_label} | 建議"):
                continue

            if " | 建議 | 建議" in normalized_line and " = " in normalized_line:
                detailed_items.append(normalized_line.split(" = ", 1)[1].strip())
                continue

            if " | 建議 = " in normalized_line:
                general_item = normalized_line.split(" | 建議 = ", 1)[1].strip()

        if detailed_items:
            return self._merge_questionnaire_recommendation_items(detailed_items)
        if general_item and general_item != "無":
            return [self._clean_questionnaire_recommendation_text(general_item)]
        return []

    def _merge_questionnaire_recommendation_items(self, items: list[str]) -> list[str]:
        merged: list[str] = []
        for item in items:
            cleaned = self._clean_questionnaire_recommendation_text(item)
            if not cleaned or cleaned == "無":
                continue

            if merged and self._is_recommendation_continuation(cleaned):
                merged[-1] = f"{merged[-1]}\n{cleaned}"
            else:
                merged.append(cleaned)
        return merged

    def _clean_questionnaire_recommendation_text(self, text: str) -> str:
        cleaned = self._normalize_text(text)
        cleaned = re.sub(r"^\d+\.\s*", "", cleaned)
        return cleaned.strip()

    def _is_recommendation_continuation(self, text: str) -> bool:
        return text.startswith(("A:", "A：", "B:", "B：", "C:", "C：", "故", "因此"))

    def _build_grounded_recommendation_sections(
        self,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> dict[str, list[str]]:
        grouped: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }

        for finding in report.key_findings:
            if not finding.description.strip() or not finding.evidence:
                continue
            title = self._classify_finding_heading(finding.item, finding.description)
            paragraph = self._build_grounded_finding_paragraph(finding.item, finding.description, finding.evidence)
            if paragraph and paragraph not in grouped[title]:
                grouped[title].append(paragraph)

        for title, paragraphs in self._build_table_grounded_recommendation_sections(normalized_document).items():
            for paragraph in paragraphs:
                if paragraph not in grouped[title]:
                    grouped[title].append(paragraph)

        for title, paragraphs in self._build_docx_grounded_recommendation_sections(normalized_document).items():
            for paragraph in paragraphs:
                if paragraph not in grouped[title]:
                    grouped[title].append(paragraph)

        return grouped

    def _build_docx_grounded_recommendation_sections(
        self,
        normalized_document: NormalizedDocument,
    ) -> dict[str, list[str]]:
        grouped: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }

        for source in normalized_document.sources:
            if source.type != "docx":
                continue

            source_label = self._humanize_source_label(source)
            for table in source.tables:
                if table.table_kind not in {"questionnaire_matrix", "diagnostic_matrix", "narrative_matrix"}:
                    continue

                semantic_lines = self._select_recommendation_semantic_lines(table.semantic_lines)
                for line in semantic_lines:
                    section_title = self._classify_semantic_line_section(line)
                    if section_title is None:
                        continue

                    paragraph = self._build_semantic_line_recommendation_paragraph(
                        source_label=source_label,
                        semantic_line=line,
                    )
                    if paragraph and paragraph not in grouped[section_title]:
                        grouped[section_title].append(paragraph)

        return grouped

    def _select_recommendation_semantic_lines(self, semantic_lines: list[str]) -> list[str]:
        detailed_lines = [
            line
            for line in semantic_lines
            if self._is_detailed_recommendation_semantic_line(line)
        ]
        if detailed_lines:
            return detailed_lines

        recommendation_lines = [
            line
            for line in semantic_lines
            if self._is_recommendation_semantic_line(line)
        ]
        return recommendation_lines

    def _is_recommendation_semantic_line(self, semantic_line: str) -> bool:
        normalized = self._normalize_text(semantic_line)
        return "建議" in normalized

    def _is_detailed_recommendation_semantic_line(self, semantic_line: str) -> bool:
        normalized = self._normalize_text(semantic_line)
        return " | 建議 | 建議" in normalized

    def _build_grounded_finding_paragraph(
        self,
        item: str,
        description: str,
        evidence: list[EvidenceReference],
    ) -> str:
        cleaned_description = description.strip().rstrip("。；")
        cleaned_evidence = self._format_evidence_references(evidence).strip().rstrip("。；")
        if not cleaned_description or not cleaned_evidence:
            return ""
        return f"依據{cleaned_evidence}，{cleaned_description}。"

    def _build_table_grounded_recommendation_sections(
        self,
        normalized_document: NormalizedDocument,
    ) -> dict[str, list[str]]:
        grouped: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }

        for source in normalized_document.sources:
            if source.type != "table":
                continue
            source_label = self._humanize_source_label(source)
            for table in source.tables:
                rows = [row for row in table.rows if isinstance(row, dict)]
                if not rows:
                    continue

                for row in rows:
                    section_title, category_label = self._map_table_category_to_section(
                        self._stringify_value(row.get("類別"))
                    )
                    if section_title is None:
                        continue

                    before_value = self._find_row_value(row, "補助前")
                    after_value = self._find_row_value(row, "補助後")
                    delta_value = self._find_row_value(row, "增減")
                    attainment_value = self._find_row_value(row, "達標")
                    focus_value = self._find_row_value(row, "補助重點")
                    benefit_value = self._find_row_value(row, "補助效益")

                    metric_parts: list[str] = []
                    if before_value:
                        metric_parts.append(f"補助前為{before_value}")
                    if after_value:
                        metric_parts.append(f"補助後為{after_value}")
                    if delta_value:
                        metric_parts.append(f"增減變化為{delta_value}")
                    if attainment_value:
                        metric_parts.append(f"孤島效應達標欄位為{attainment_value}")
                    if metric_parts:
                        grouped[section_title].append(
                            f"依據{source_label}，{category_label}項目{self._join_metric_parts(metric_parts)}。"
                        )

                    focus_text = self._sanitize_table_note(focus_value)
                    benefit_text = self._sanitize_table_note(benefit_value)
                    note_parts: list[str] = []
                    if focus_text:
                        note_parts.append(f"補助重點為{focus_text}")
                    if benefit_text:
                        note_parts.append(f"補助效益為{benefit_text}")
                    if note_parts:
                        grouped[section_title].append(
                            f"同一{source_label}顯示，{category_label}項目{self._join_metric_parts(note_parts)}。"
                        )

        return grouped

    def _classify_semantic_line_section(self, semantic_line: str) -> str | None:
        normalized_title = self._normalize_recommendation_section_title(semantic_line)
        supported_titles = {
            "供電部分",
            "供水部分",
            "供氧氣及IT系統部分之建議",
            "其他(工安衛)",
        }
        return normalized_title if normalized_title in supported_titles else None

    def _build_semantic_line_recommendation_paragraph(
        self,
        *,
        source_label: str,
        semantic_line: str,
    ) -> str:
        normalized_line = self._normalize_text(semantic_line)
        if " = " not in normalized_line:
            return ""

        prefix, value = normalized_line.split(" = ", 1)
        cleaned_value = self._sanitize_table_note(value)
        if not cleaned_value:
            return ""

        if any(
            keyword in prefix
            for keyword in ("醫院名稱", "主辦單位", "日期", "診斷專家", "執行機構", "執行期間")
        ):
            return ""

        readable_prefix = prefix.replace("診斷表 | ", "").replace("問卷欄位 | ", "").replace("問卷敘述 | ", "")
        readable_prefix = readable_prefix.replace("基本資料 | ", "").strip(" |")
        if not readable_prefix:
            return ""

        return f"依據{source_label}，{readable_prefix}為{cleaned_value}。"

    def _map_table_category_to_section(self, raw_category: str) -> tuple[str | None, str]:
        category = raw_category.strip().lower()
        if category in {"電", "供電"}:
            return "供電部分", "供電"
        if category in {"水", "供水"}:
            return "供水部分", "供水"
        if category in {"o2", "氧", "供氧", "供氣"}:
            return "供氧氣及IT系統部分之建議", "供氧"
        if category in {"it", "資訊", "資訊系統", "it備援"}:
            return "供氧氣及IT系統部分之建議", "IT備援"
        return None, raw_category.strip() or "該項目"

    def _find_row_value(self, row: dict[str, object], keyword: str) -> str:
        for key, value in row.items():
            if keyword in key:
                text = self._stringify_value(value)
                if text:
                    return text
        return ""

    def _sanitize_table_note(self, value: str) -> str:
        text = self._stringify_value(value)
        if not text:
            return ""
        sanitized = re.sub(r"\s+", " ", text.replace("\n", "；")).strip("； ")
        return sanitized[:140] if len(sanitized) > 140 else sanitized

    def _join_metric_parts(self, parts: list[str]) -> str:
        cleaned_parts = [part.strip().rstrip("。；") for part in parts if part.strip()]
        if not cleaned_parts:
            return ""
        if len(cleaned_parts) == 1:
            return cleaned_parts[0]
        return "，".join(cleaned_parts)

    def _humanize_source_label(self, source: ExtractedSource) -> str:
        source_role = self._classify_source_role(source)
        if source_role == "benefit_table":
            return "效益調查簡表"
        if source_role == "questionnaire_recommendation_docx":
            return "問卷診斷建議單"
        if source_role == "diagnostic_docx":
            return "專家診斷表"
        if source_role == "questionnaire_docx":
            return "問卷資料"
        return f"「{Path(source.source_name).stem}」"

    def _classify_source_role(self, source: ExtractedSource) -> str:
        if source.type == "table":
            return "benefit_table" if self._is_benefit_table_source(source) else "generic_table"

        if source.type != "docx":
            return "generic"

        if self._has_questionnaire_recommendation_content(source):
            return "questionnaire_recommendation_docx"

        table_kinds = {table.table_kind for table in source.tables}
        if "diagnostic_matrix" in table_kinds:
            return "diagnostic_docx"
        if table_kinds & {"questionnaire_matrix", "metadata_block", "narrative_matrix"}:
            return "questionnaire_docx"
        return "generic_docx"

    def _is_benefit_table_source(self, source: ExtractedSource) -> bool:
        for table in source.tables:
            rows = [row for row in table.rows if isinstance(row, dict)]
            if not rows:
                continue

            row_keys = {self._normalize_text(str(key)) for row in rows for key in row.keys()}
            if (
                any("類別" in key for key in row_keys)
                and any("補助前" in key for key in row_keys)
                and any("補助後" in key for key in row_keys)
                and any("補助效益" in key or "補助重點" in key for key in row_keys)
            ):
                return True
        return False

    def _build_recommendation_sections(self, report: ReportJSON) -> list[ReportSection]:
        grouped: dict[str, list[str]] = {
            "供電部分": [],
            "供水部分": [],
            "供氧氣及IT系統部分之建議": [],
            "其他(工安衛)": [],
        }

        for finding in report.key_findings:
            if not finding.evidence:
                continue
            title = self._classify_finding_heading(finding.item, finding.description)
            grouped[title].append(finding.description)

        sections: list[ReportSection] = []
        for title, paragraphs in grouped.items():
            cleaned = self._clean_text_items(paragraphs)
            if cleaned:
                sections.append(ReportSection(title=title, paragraphs=cleaned))

        remaining_recommendations = [
            item
            for item in self._evidenced_text_list_to_strings(
                report.recommendations,
                require_evidence=True,
            )
            if item
            and item not in {paragraph for paragraphs in grouped.values() for paragraph in paragraphs}
        ]
        if remaining_recommendations:
            sections.append(ReportSection(title="補充建議", paragraphs=remaining_recommendations))

        return sections

    def _evidenced_text_to_string(
        self,
        item: EvidenceText,
        *,
        require_evidence: bool,
    ) -> str:
        if require_evidence and not self._has_evidence(item):
            return ""
        return item.text.strip()

    def _evidenced_text_list_to_strings(
        self,
        items: list[EvidenceText],
        *,
        require_evidence: bool,
    ) -> list[str]:
        return self._clean_text_items(
            [
                text
                for item in items
                if (
                    text := self._evidenced_text_to_string(
                        item,
                        require_evidence=require_evidence,
                    )
                )
            ]
        )

    def _has_evidence(self, item: EvidenceText) -> bool:
        return any(
            evidence.source_name.strip() or evidence.content.strip()
            for evidence in item.evidence
        )

    def _format_evidence_references(self, evidence_items: list[EvidenceReference]) -> str:
        parts: list[str] = []
        for evidence in evidence_items[:3]:
            source = evidence.source_name.strip() or "來源資料"
            content = self._normalize_text(evidence.content)
            if len(content) > 120:
                content = f"{content[:117]}..."
            if content:
                parts.append(f"{source}（{evidence.fact_type}：{content}）")
            else:
                parts.append(source)
        return "；".join(parts)

    def _normalize_recommendation_section_title(self, title: str) -> str:
        normalized = self._normalize_text(title)
        if "供電" in normalized or "發電" in normalized or "油" in normalized:
            return "供電部分"
        if "供水" in normalized or "蓄水" in normalized or "水箱" in normalized:
            return "供水部分"
        if "供氧" in normalized or "氧氣" in normalized or "IT" in normalized or "資訊" in normalized:
            return "供氧氣及IT系統部分之建議"
        if "工安衛" in normalized or "其他" in normalized:
            return "其他(工安衛)"
        return title.strip()

    def _collect_analysis_chart_sources(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[ExtractedSource]:
        return [source for source in normalized_document.sources if source.type == "image"]

    def _format_numbered_section_title(self, index: int, title: str) -> str:
        clean_title = title.strip()
        if re.match(r"^[一二三四五六七八九十]+、", clean_title):
            return clean_title
        return f"{self._to_chinese_numeral(index)}、{clean_title}"

    def _format_numbered_subitem(self, index: int, text: str) -> str:
        clean_text = text.strip()
        if re.match(r"^（[一二三四五六七八九十]+）", clean_text):
            return clean_text
        return f"（{self._to_chinese_numeral(index)}） {clean_text}"

    def _to_chinese_numeral(self, value: int) -> str:
        numerals = {
            1: "一",
            2: "二",
            3: "三",
            4: "四",
            5: "五",
            6: "六",
            7: "七",
            8: "八",
            9: "九",
            10: "十",
        }
        if value in numerals:
            return numerals[value]
        if value < 20:
            return f"十{numerals[value - 10]}"
        tens, ones = divmod(value, 10)
        if ones == 0:
            return f"{numerals.get(tens, str(tens))}十"
        return f"{numerals.get(tens, str(tens))}十{numerals.get(ones, str(ones))}"

    def _build_diagnosis_tables(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[ReportTable]:
        diagnosis_tables: list[ReportTable] = []

        for source in normalized_document.sources:
            if source.type != "table" or self._is_benefit_table_source(source):
                continue
            for extracted_table in source.tables[:1]:
                headers, rows = self._normalize_extracted_table(extracted_table)
                if not headers and not rows:
                    continue
                if not headers and rows:
                    headers = [f"欄位{index + 1}" for index in range(len(rows[0]))]

                compact_headers, compact_rows = self._compact_report_table(headers, rows[:8])
                if not compact_headers or not compact_rows:
                    continue

                diagnosis_tables.append(
                    ReportTable(
                        title=self._normalize_analysis_table_title(
                            extracted_table.sheet_name or extracted_table.title or source.source_name
                        ),
                        columns=compact_headers,
                        rows=compact_rows,
                    )
                )
                return diagnosis_tables

        return diagnosis_tables

    def _build_analysis_tables(
        self,
        report: ReportJSON,
        normalized_document: NormalizedDocument,
    ) -> list[ReportTable]:
        analysis_tables: list[ReportTable] = []

        for table in report.tables:
            analysis_tables.append(
                table.model_copy(update={"title": self._normalize_analysis_table_title(table.title)})
            )

        if analysis_tables:
            return analysis_tables

        for source in normalized_document.sources:
            if source.type != "table":
                continue
            for extracted_table in source.tables[:1]:
                headers, rows = self._normalize_extracted_table(extracted_table)
                if not headers and not rows:
                    continue
                if not headers and rows:
                    headers = [f"欄位{index + 1}" for index in range(len(rows[0]))]
                analysis_tables.append(
                    ReportTable(
                        title=self._normalize_analysis_table_title(
                            extracted_table.sheet_name or extracted_table.title or source.source_name
                        ),
                        columns=headers,
                        rows=rows[:8],
                    )
                )
                return analysis_tables

        return analysis_tables

    def _compact_report_table(
        self,
        headers: list[str],
        rows: list[list[str]],
    ) -> tuple[list[str], list[list[str]]]:
        if not headers:
            return headers, rows

        kept_indexes: list[int] = []
        for index, header in enumerate(headers):
            column_values = [
                row[index].strip()
                for row in rows
                if index < len(row) and row[index] is not None
            ]
            if any(column_values):
                kept_indexes.append(index)

        compact_headers = [headers[index] for index in kept_indexes]
        compact_rows = [
            [row[index] if index < len(row) else "" for index in kept_indexes]
            for row in rows
        ]
        return compact_headers, compact_rows

    def _classify_finding_heading(self, item: str, description: str) -> str:
        combined = f"{item} {description}"
        if any(keyword in combined for keyword in ("供電", "發電", "UPS", "油槽", "油量")):
            return "供電部分"
        if any(keyword in combined for keyword in ("供水", "蓄水", "水塔", "水箱")):
            return "供水部分"
        if any(keyword in combined for keyword in ("供氧", "氧氣", "IT", "資訊", "機房")):
            return "供氧氣及IT系統部分之建議"
        return "其他(工安衛)"

    def _render_report_tables(self, document: Document, tables: list[ReportTable]) -> None:
        for report_table in tables:
            self._append_paragraph(
                document,
                self._normalize_analysis_table_title(report_table.title),
                style_name="Normal",
                bold=True,
            )
            headers = report_table.columns or ["欄位"]
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header

            for row in report_table.rows:
                row_cells = table.add_row().cells
                normalized_row = row + [""] * (len(headers) - len(row))
                for index, value in enumerate(normalized_row[: len(headers)]):
                    row_cells[index].text = value

    def _normalize_extracted_table(
        self,
        extracted_table: ExtractedTable,
    ) -> tuple[list[str], list[list[str]]]:
        raw_rows = extracted_table.rows
        if not raw_rows:
            return extracted_table.columns, []

        if isinstance(raw_rows[0], dict):
            headers = extracted_table.columns or list(raw_rows[0].keys())
            rows = [
                [self._stringify_value(row.get(header)) for header in headers]
                for row in raw_rows
            ]
            return headers, rows

        rows_as_lists = [
            [self._stringify_value(cell) for cell in row]
            for row in raw_rows
            if isinstance(row, list)
        ]
        headers = extracted_table.columns
        if not headers and rows_as_lists:
            headers = rows_as_lists[0]
            rows_as_lists = rows_as_lists[1:]
        elif headers and rows_as_lists and rows_as_lists[0] == headers:
            rows_as_lists = rows_as_lists[1:]

        return headers, rows_as_lists

    def _normalize_attachment_table_title(
        self,
        title: str | None,
        source_name: str,
        fallback: str,
    ) -> str:
        candidate = (title or "").strip()
        if candidate and not self._is_generic_table_title(candidate):
            return candidate

        source_stem = Path(source_name).stem
        if source_stem:
            return source_stem
        return fallback

    def _normalize_analysis_table_title(self, title: str) -> str:
        candidate = title.strip() if title else ""
        if not candidate or self._is_generic_table_title(candidate):
            return "申請韌性補助計畫導入前後效益調查簡表"
        return candidate

    def _is_generic_table_title(self, title: str) -> bool:
        normalized = self._normalize_text(title).lower()
        return bool(
            re.fullmatch(r"(sheet|worksheet)\s*\d+", normalized)
            or re.fullmatch(r"工作表\s*\d+", normalized)
            or re.fullmatch(r"table\s*\d+", normalized)
        )

    def _add_title(self, document: Document, title: str, template_name: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.style = "Title"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(f"模板：{template_name}")

    def _add_heading(self, document: Document, heading: str) -> Paragraph:
        return document.add_heading(heading, level=1)

    def _append_paragraph(
        self,
        document: Document,
        text: str,
        style_name: str | None = None,
        bold: bool = False,
    ) -> None:
        paragraph = document.add_paragraph()
        if style_name:
            paragraph.style = style_name
        run = paragraph.add_run(text)
        run.bold = bold
        self._apply_run_font(run)

    def _add_paragraphs(self, document: Document, paragraphs: list[str]) -> None:
        if not paragraphs:
            self._append_paragraph(document, "資料不足", style_name=self.LIST_STYLE)
            return

        for paragraph in paragraphs:
            text = paragraph.strip()
            if not text:
                continue
            self._append_paragraph(document, text, style_name=self.LIST_STYLE)

    def _add_bullets(self, document: Document, items: list[str]) -> None:
        if not items:
            self._append_paragraph(document, "資料不足", style_name=self.BULLET_STYLE)
            return

        for item in items:
            self._append_paragraph(document, item, style_name=self.BULLET_STYLE)

    def _load_source_docx(self, source: ExtractedSource) -> Document | None:
        source_path_value = source.metadata.get("path")
        if not source_path_value:
            return None

        source_path = Path(source_path_value)
        if not source_path.is_file():
            return None

        return Document(source_path)

    def _build_fixed_introduction_paragraphs(self, hospital_name: str) -> list[str]:
        resolved_hospital_name = hospital_name.strip() or "OOOO醫院"
        return [
            (
                "因應國際局勢變化，國際情勢、極端氣候、人為或自然災害等重大事故或特殊事件，"
                "我國醫療體系亟須強化醫院設備韌性應變能力，以維繫院內維生系統基本運作，"
                "促使醫院有較長醫療之黃金時間，維繫國民生命安全與健康。"
            ),
            (
                "爰本報告係根據本中心執行衛福部114-115年度「急救責任醫院維生系統韌性盤點暨輔導計畫」"
                "執行內容之輔導面分項中所列『配合衛福部「114-115 年提升關鍵基礎設施醫院維生系統韌性補助計畫」，"
                "提出 56 家重要急救責任醫院可行之維生系統設備韌性備援能力診斷輔導方案報告』，"
                f"針對{resolved_hospital_name}就其供電(供油量)、供水、供氣(氧氣)及資訊系統等項目，"
                "診斷提出孤島效應下提升備援能力之建議，以供參考。"
            ),
        ]

    def _build_chapter_2_context(
        self,
        hospital_name: str,
        normalized_document: NormalizedDocument,
    ) -> dict[str, str | bool]:
        hospital_label = hospital_name.strip() or "OOOO醫院"
        diagnosis_hospital_name = self._with_ministry_prefix(hospital_label)
        return {
            "hospital_name": hospital_label,
            "diagnosis_hospital_name": diagnosis_hospital_name,
            "has_visit_image": self._collect_diagnosis_image_source(normalized_document) is not None,
            "visit_date": self._infer_visit_date(normalized_document) or "",
            "overall_capacity_phrase": self._infer_diagnosis_overall_capacity_phrase(normalized_document),
            "generator_risk_phrase": self._infer_diagnosis_generator_risk_phrase(normalized_document),
            "it_risk_phrase": self._infer_diagnosis_it_risk_phrase(normalized_document),
            "subsidy_status_phrase": self._infer_diagnosis_subsidy_status_phrase(normalized_document),
        }

    def _build_diagnosis_lead_paragraph(
        self,
        chapter_2_context: dict[str, str | bool],
    ) -> str:
        diagnosis_hospital_name = str(chapter_2_context["diagnosis_hospital_name"])
        if chapter_2_context.get("has_visit_image"):
            return (
                f"根據訪視診斷{diagnosis_hospital_name}(如下圖)，該院在供電(供油量)、供水、"
                "供氣(氧氣)及資訊系統等項目運作量能表現及總體建議如問卷填答內容(如附表)。"
            )
        return (
            f"根據訪視診斷{diagnosis_hospital_name}，該院在供電(供油量)、供水、"
            "供氣(氧氣)及資訊系統等項目運作量能表現及總體建議如問卷填答內容(如附表)。"
        )

    def _build_fixed_diagnosis_paragraphs(
        self,
        chapter_2_context: dict[str, str | bool],
    ) -> list[str]:
        resolved_hospital_name = str(chapter_2_context["diagnosis_hospital_name"])
        overall_capacity_phrase = str(chapter_2_context["overall_capacity_phrase"])
        generator_risk_phrase = str(chapter_2_context["generator_risk_phrase"])
        it_risk_phrase = str(chapter_2_context["it_risk_phrase"])
        subsidy_status_phrase = str(chapter_2_context["subsidy_status_phrase"])

        observation_paragraph = (
            f"根據診斷總體觀察，{resolved_hospital_name}在供電(供油量)、供水、供氣(氧氣)等項目"
            f"考量孤島效應下{overall_capacity_phrase}，惟{generator_risk_phrase}，另外{it_risk_phrase}。"
        )

        subsidy_paragraph = (
            f"爰，{subsidy_status_phrase}，未來可參考下節專家診斷建議評估"
            "結合韌性資源提高備援運作量能。"
        )

        return [observation_paragraph, subsidy_paragraph]

    def _has_table_kind(
        self,
        normalized_document: NormalizedDocument,
        target_kinds: set[str],
    ) -> bool:
        for source in normalized_document.sources:
            for table in source.tables:
                if table.table_kind in target_kinds:
                    return True
        return False

    def _extract_benefit_categories(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[str]:
        categories: list[str] = []
        seen: set[str] = set()
        for source in normalized_document.sources:
            if source.type != "table":
                continue
            for table in source.tables:
                for row in table.rows:
                    if not isinstance(row, dict):
                        continue
                    for key, value in row.items():
                        if "類別" not in self._normalize_text(str(key)):
                            continue
                        category = self._normalize_text(self._stringify_value(value))
                        if not category or category in seen:
                            continue
                        seen.add(category)
                        categories.append(category)
        return categories[:6]

    def _infer_diagnosis_overall_capacity_phrase(
        self,
        normalized_document: NormalizedDocument,
    ) -> str:
        power_ok = self._domain_reaches_72_hours(normalized_document, ("供電", "發電機", "UPS", "柴油", "油槽"))
        water_ok = self._domain_reaches_72_hours(normalized_document, ("供水", "用水", "蓄水", "儲水", "水塔"))
        oxygen_ok = self._domain_reaches_72_hours(normalized_document, ("供氣", "氧氣", "液氧", "氧"))

        positive_count = sum(1 for status in (power_ok, water_ok, oxygen_ok) if status is True)
        unknown_count = sum(1 for status in (power_ok, water_ok, oxygen_ok) if status is None)

        if positive_count == 3:
            return "各運作量能均可能達72小時"
        if positive_count >= 2 and unknown_count <= 1:
            return "多數項目運作量能可能達72小時，惟部分項目仍待進一步確認"
        if positive_count >= 1:
            return "部分項目運作量能可能達72小時，惟整體仍需持續盤點"
        return "各運作量能是否均可達72小時仍待進一步確認"

    def _infer_diagnosis_generator_risk_phrase(
        self,
        normalized_document: NormalizedDocument,
    ) -> str:
        texts = self._collect_questionnaire_analysis_texts(
            normalized_document,
            ("供電", "發電機", "柴油", "油槽"),
        )
        if self._contains_any(texts, ("僅有一台發電機", "只有一台發電機", "1台發電機", "一台發電機")):
            return "發電機僅一台，在面對不可預知的重大災難或緊急事故，存在一定風險"
        if self._contains_any(texts, ("兩台發電機", "2台發電機", "雙機", "雙台發電機")):
            return "發電機雖具多機配置，仍應持續注意設備維護與故障切換風險"
        return "發電機配置與故障切換風險仍待依問卷資料持續確認"

    def _infer_diagnosis_it_risk_phrase(
        self,
        normalized_document: NormalizedDocument,
    ) -> str:
        texts = self._collect_questionnaire_analysis_texts(
            normalized_document,
            ("資訊系統", "IT", "異地", "異機", "單機", "備援機制", "網路"),
        )
        lacks_multi_backup = self._contains_any(
            texts,
            ("缺乏異機", "缺乏異地", "無異地", "無異機", "僅能單機", "只有單機", "單機作業"),
        )
        has_multi_backup = self._contains_any(
            texts,
            ("異地備援", "異機備援", "多元備援", "雙機備援", "跨站備援"),
        )

        if lacks_multi_backup and not has_multi_backup:
            return "資訊系統備援部分除能單機作業外，缺乏異機或異地等多元備援方案"
        if has_multi_backup:
            return "資訊系統已具部分備援機制，但仍應持續盤點通訊維持與切換演練能力"
        return "資訊系統備援機制與異地異機配置情形仍待進一步確認"

    def _infer_diagnosis_subsidy_status_phrase(
        self,
        normalized_document: NormalizedDocument,
    ) -> str:
        has_benefit_table = any(
            source.type == "table" and self._is_benefit_table_source(source)
            for source in normalized_document.sources
        )
        texts = self._collect_questionnaire_analysis_texts(
            normalized_document,
            ("補助", "申請", "衛福部"),
        )

        if has_benefit_table:
            return (
                "依現有資料，補助計畫申請與導入狀態仍需再確認；"
                "惟效益調查簡表已呈現補助前後韌性天數變化（如下表）"
            )

        if self._contains_any(texts, ("■申請衛福部", "已申請補助計畫", "曾申請補助計畫", "已申請", "曾申請")):
            return "依現有資料，該院已申請補助計畫"
        if self._contains_any(texts, ("擬申請", "預計申請", "研議申請", "未來將作整體評估")):
            return "依現有資料，該院已規劃或研議申請衛福部韌性補助計畫"
        if self._contains_any(texts, ("未申請補助計畫", "尚未申請補助計畫", "尚未申請", "未申請")):
            return "依現有資料，該院補助計畫申請情形為尚未申請"
        return "該院補助計畫申請情形依現有資料尚無法完整判定"

    def _domain_reaches_72_hours(
        self,
        normalized_document: NormalizedDocument,
        keywords: tuple[str, ...],
    ) -> bool | None:
        texts = self._collect_questionnaire_analysis_texts(normalized_document, keywords)
        if not texts:
            return None

        for text in texts:
            normalized = self._normalize_text(text).lower()
            if re.search(r"(超過|達|可達|達到)\s*72\s*(小時|hr|hrs)", normalized):
                return True
            if re.search(r"72\s*(小時|hr|hrs)\s*以上", normalized):
                return True
            if re.search(r"(\d+(?:\.\d+)?)\s*(天|日)", normalized):
                day_values = [
                    float(match.group(1))
                    for match in re.finditer(r"(\d+(?:\.\d+)?)\s*(天|日)", normalized)
                ]
                if any(value >= 3 for value in day_values):
                    return True
            if re.search(r"(\d+(?:\.\d+)?)\s*(小時|hr|hrs)", normalized):
                hour_values = [
                    float(match.group(1))
                    for match in re.finditer(r"(\d+(?:\.\d+)?)\s*(小時|hr|hrs)", normalized)
                ]
                if any(value >= 72 for value in hour_values):
                    return True

        for text in texts:
            normalized = self._normalize_text(text)
            if any(negative_phrase in normalized for negative_phrase in ("未達72小時", "不足72小時", "低於72小時")):
                return False
        return None

    def _collect_questionnaire_analysis_texts(
        self,
        normalized_document: NormalizedDocument,
        keywords: tuple[str, ...],
    ) -> list[str]:
        questionnaire_only = NormalizedDocument(
            project_name=normalized_document.project_name,
            hospital_name=normalized_document.hospital_name,
            template_id=normalized_document.template_id,
            sources=[
                source
                for source in normalized_document.sources
                if self._classify_source_role(source) in {"questionnaire_docx", "questionnaire_recommendation_docx"}
            ],
            facts=[
                fact
                for fact in normalized_document.facts
                if any(
                    self._normalize_text(source.source_name) == self._normalize_text(fact.source_name)
                    for source in normalized_document.sources
                    if self._classify_source_role(source) in {"questionnaire_docx", "questionnaire_recommendation_docx"}
                )
            ],
            warnings=normalized_document.warnings,
        )
        return self._collect_analysis_texts(questionnaire_only, keywords)

    def _collect_analysis_texts(
        self,
        normalized_document: NormalizedDocument,
        keywords: tuple[str, ...],
    ) -> list[str]:
        matches: list[str] = []
        seen: set[str] = set()

        def add_candidate(text: str) -> None:
            normalized = self._normalize_text(text)
            if not normalized:
                return
            lowered = normalized.lower()
            if not any(keyword.lower() in lowered for keyword in keywords):
                return
            if normalized in seen:
                return
            seen.add(normalized)
            matches.append(normalized)

        for fact in normalized_document.facts:
            add_candidate(fact.content)

        for source in normalized_document.sources:
            add_candidate(source.source_name)
            for paragraph in source.paragraphs:
                add_candidate(paragraph)
            if source.text:
                add_candidate(source.text)
            for table in source.tables:
                for semantic_line in table.semantic_lines:
                    add_candidate(semantic_line)
                for cell in table.cells:
                    add_candidate(cell.text)
                for row in table.rows:
                    if isinstance(row, dict):
                        add_candidate("；".join(f"{key}: {value}" for key, value in row.items() if value not in (None, "")))
                    elif isinstance(row, list):
                        add_candidate("；".join(self._stringify_value(value) for value in row if value not in (None, "")))

        return matches

    def _contains_any(self, texts: list[str], phrases: tuple[str, ...]) -> bool:
        return any(phrase in text for text in texts for phrase in phrases)

    def _collect_diagnosis_image_source(
        self,
        normalized_document: NormalizedDocument,
    ) -> ExtractedSource | None:
        for source in normalized_document.sources:
            if source.type == "image":
                return source
        return None

    def _build_diagnosis_image_caption(
        self,
        chapter_2_context: dict[str, str | bool],
    ) -> str:
        visit_date = str(chapter_2_context.get("visit_date", "")).strip()
        hospital_label = str(chapter_2_context.get("hospital_name", "OOOO醫院")).strip() or "OOOO醫院"
        prefix = f"{visit_date}" if visit_date else ""
        return f"{prefix}前往{hospital_label}進行設備韌性備援能力診斷會議中留影"

    def _infer_visit_date(self, normalized_document: NormalizedDocument) -> str | None:
        gregorian_pattern = re.compile(r"(20\d{2})[./-](\d{1,2})[./-](\d{1,2})")
        roc_pattern = re.compile(r"民國\s*(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日")

        text_candidates: list[str] = []
        for source in normalized_document.sources:
            text_candidates.append(source.source_name)
            text_candidates.extend(source.paragraphs)
            if source.text:
                text_candidates.append(source.text)
            for table in source.tables:
                for cell in table.cells:
                    text_candidates.append(cell.text)

        for text in text_candidates:
            gregorian_match = gregorian_pattern.search(text)
            if gregorian_match:
                year, month, day = (int(part) for part in gregorian_match.groups())
                return f"{year:04d}.{month:02d}.{day:02d}"

        for text in text_candidates:
            roc_match = roc_pattern.search(text)
            if roc_match:
                roc_year, month, day = (int(part) for part in roc_match.groups())
                return f"{roc_year + 1911:04d}.{month:02d}.{day:02d}"

        return None

    def _with_ministry_prefix(self, hospital_name: str) -> str:
        normalized_hospital_name = hospital_name.strip() or "OOOO醫院"
        if normalized_hospital_name.startswith("衛生福利部"):
            return normalized_hospital_name
        return f"衛生福利部{normalized_hospital_name}"

    def _fallback_introduction_paragraphs(self, report: ReportJSON) -> list[str]:
        candidates = report.background.strip()
        if candidates:
            return [candidates]
        return ["資料不足"]

    def _fallback_diagnosis_paragraphs(self, report: ReportJSON) -> list[str]:
        paragraphs: list[str] = []
        if report.summary.strip():
            paragraphs.append(report.summary.strip())
        if report.key_findings:
            paragraphs.extend(
                finding.description
                for finding in report.key_findings[:3]
                if finding.evidence
            )
        return paragraphs or ["資料不足"]

    def _infer_hospital_name(
        self,
        project_name: str,
        normalized_document: NormalizedDocument,
    ) -> str:
        explicit_hospital_name = (normalized_document.hospital_name or "").strip()
        if explicit_hospital_name:
            return explicit_hospital_name

        clean_project_name = project_name.strip()
        if self._looks_like_hospital_name(clean_project_name):
            return clean_project_name

        explicit_name_pattern = re.compile(r"醫院名稱[:：]\s*([^\n，。；（）()]{1,30}醫院)")
        ministry_name_pattern = re.compile(r"(衛生福利部[^\s，。；（）()]{1,24}醫院)")
        generic_name_pattern = re.compile(r"([^\s，。；（）()]{1,24}醫院)")

        explicit_candidates: list[str] = []
        ministry_candidates: list[str] = []
        generic_candidates: list[str] = []

        def collect_from_text(text: str) -> None:
            explicit_candidates.extend(explicit_name_pattern.findall(text))
            ministry_candidates.extend(ministry_name_pattern.findall(text))
            for candidate in generic_name_pattern.findall(text):
                candidate = candidate.strip()
                if not self._looks_like_hospital_name(candidate):
                    continue
                generic_candidates.append(candidate)

        for source in normalized_document.sources:
            for paragraph in source.paragraphs:
                collect_from_text(paragraph)
            if source.text:
                collect_from_text(source.text)
            for table in source.tables:
                for cell in table.cells:
                    collect_from_text(cell.text)

        for candidates in (explicit_candidates, ministry_candidates, generic_candidates):
            cleaned_candidates = [candidate.strip() for candidate in candidates if candidate.strip()]
            if cleaned_candidates:
                return cleaned_candidates[0]

        return clean_project_name or "OO醫院"

    def _looks_like_hospital_name(self, value: str) -> bool:
        if not value or "醫院" not in value:
            return False

        generic_keywords = (
            "報告",
            "模板",
            "範本",
            "診斷",
            "輔導",
            "附件",
            "計畫",
            "生成",
        )
        if any(keyword in value and "醫院" not in keyword for keyword in generic_keywords):
            if value.endswith("醫院"):
                return True
            return False

        return True

    def _is_benefit_table_source(self, source: ExtractedSource) -> bool:
        for table in source.tables:
            rows = [row for row in table.rows if isinstance(row, dict)]
            if not rows:
                continue

            row_keys = {self._normalize_text(str(key)) for row in rows for key in row.keys()}
            if (
                any("類別" in key for key in row_keys)
                and any("補助前" in key for key in row_keys)
                and any("補助後" in key for key in row_keys)
                and any("補助效益" in key or "補助重點" in key for key in row_keys)
            ):
                return True
        return False

    def _build_diagnosis_benefit_summary_rows(
        self,
        normalized_document: NormalizedDocument,
    ) -> list[list[str]]:
        for source in normalized_document.sources:
            if source.type != "table":
                continue

            for table in source.tables:
                rows = [row for row in table.rows if isinstance(row, dict)]
                if not rows:
                    continue

                row_keys = {self._normalize_text(str(key)) for row in rows for key in row.keys()}
                if not (
                    any("類別" in key for key in row_keys)
                    and any("補助前" in key for key in row_keys)
                    and any("補助後" in key for key in row_keys)
                ):
                    continue

                summary_rows: list[list[str]] = []
                for row in rows:
                    category = self._stringify_value(self._find_row_value(row, "類別")).strip()
                    if not category or category.startswith("以上"):
                        continue

                    summary_rows.append(
                        [
                            category,
                            self._format_benefit_summary_value(self._find_row_value(row, "補助前")),
                            self._format_benefit_summary_value(self._find_row_value(row, "補助後")),
                            self._format_benefit_summary_value(self._find_row_value(row, "增減變化")),
                            self._format_benefit_attainment_value(self._find_row_value(row, "孤島效應達標")),
                            self._format_benefit_note_value(self._find_row_value(row, "補助重點")),
                            self._format_benefit_note_value(self._find_row_value(row, "補助效益")),
                            self._format_benefit_note_value(self._find_row_value(row, "備註")),
                        ]
                    )

                if summary_rows:
                    return summary_rows

        return []

    def _insert_diagnosis_benefit_summary_table_before(
        self,
        document: Document,
        anchor_paragraph: Paragraph,
        summary_rows: list[list[str]],
    ) -> None:
        table = document.add_table(rows=2 + len(summary_rows), cols=8)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        top_headers = ["", "天數", "天數", "增減變化天數", "孤島效應達標", "補助重點", "補助效益", "備註"]
        bottom_headers = ["類別", "補助前", "補助後", "", "", "", "", ""]

        for column_index, value in enumerate(top_headers):
            self._set_table_cell_text(
                table.rows[0].cells[column_index],
                value,
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for column_index, value in enumerate(bottom_headers):
            self._set_table_cell_text(
                table.rows[1].cells[column_index],
                value,
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for column_index in range(3, 8):
            merged_cell = table.rows[0].cells[column_index].merge(table.rows[1].cells[column_index])
            self._set_table_cell_text(
                merged_cell,
                top_headers[column_index],
                font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                bold=True,
                font_color=RGBColor(255, 255, 255),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                fill=self.DIAGNOSIS_SUMMARY_HEADER_FILL,
            )

        for row_index, row_values in enumerate(summary_rows, start=2):
            for column_index, value in enumerate(row_values):
                fill = self.DIAGNOSIS_SUMMARY_HIGHLIGHT_FILL if column_index == 1 else None
                alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index <= 4 else WD_ALIGN_PARAGRAPH.LEFT
                self._set_table_cell_text(
                    table.rows[row_index].cells[column_index],
                    value,
                    font_size=self.DIAGNOSIS_SUMMARY_TABLE_FONT_SIZE,
                    bold=False,
                    font_color=RGBColor(0, 0, 0),
                    alignment=alignment,
                    vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER,
                    fill=fill,
                )

        anchor_paragraph._p.addprevious(table._tbl)

    def _set_table_cell_text(
        self,
        cell: object,
        text: str,
        *,
        font_size: Pt,
        bold: bool,
        font_color: RGBColor,
        alignment: WD_ALIGN_PARAGRAPH,
        vertical_alignment: WD_CELL_VERTICAL_ALIGNMENT,
        fill: str | None = None,
    ) -> None:
        cell.text = text
        cell.vertical_alignment = vertical_alignment
        if fill:
            self._set_cell_fill(cell, fill)

        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                self._apply_run_font(
                    run,
                    font_name=self.BODY_FONT_NAME,
                    font_size=font_size,
                    bold=bold,
                )
                run.font.color.rgb = font_color

    def _set_cell_fill(self, cell: object, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    def _format_benefit_summary_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""

        numeric_value = self._to_float(text)
        if numeric_value is None:
            return text
        if numeric_value.is_integer():
            return str(int(numeric_value))
        return f"{numeric_value:.2f}".rstrip("0").rstrip(".")

    def _format_benefit_attainment_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""

        normalized = text.lower()
        if normalized in {"v", "y", "yes", "true", "達標"}:
            return "達標"
        if normalized in {"x", "n", "no", "false", "未達標"}:
            return "未達標"
        return text

    def _format_benefit_note_value(self, value: str) -> str:
        text = self._stringify_value(value).strip()
        if not text:
            return ""
        return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()

    def _clean_text_items(self, items: list[str]) -> list[str]:
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in items:
            text = item.strip()
            if not text:
                continue
            if text in seen:
                continue
            cleaned.append(text)
            seen.add(text)
        return cleaned

    def _format_roc_date(self) -> str:
        now = datetime.now()
        roc_year = now.year - 1911
        return f"中華民國 {roc_year} 年 {now.month} 月 {now.day} 日"

    def _split_blocks(self, text: str, limit: int) -> list[str]:
        blocks: list[str] = []
        for line in text.splitlines():
            cleaned = line.strip()
            if cleaned:
                blocks.append(cleaned)
            if len(blocks) >= limit:
                break
        return blocks or ["資料不足"]

    def _to_float(self, value: object) -> float | None:
        if value is None:
            return None
        if isinstance(value, (int, float)):
            return float(value)

        text = str(value).strip().replace(",", "")
        if not text:
            return None
        try:
            return float(text)
        except ValueError:
            return None

    def _stringify_value(self, value: object) -> str:
        if value is None:
            return ""
        return str(value).replace("\n", " ").strip()

    def _normalize_text(self, text: str) -> str:
        return " ".join(text.split()).strip()

    def _build_output_name(self, project_name: str) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = re.sub(r'[\\/:*?"<>|]+', "_", project_name).strip()
        safe_name = safe_name.replace(" ", "_") or "report"
        return f"{timestamp}_{safe_name}_{uuid4().hex[:8]}.docx"

    def _resolve_output_label(
        self,
        project_name: str,
        normalized_document: NormalizedDocument,
    ) -> str:
        clean_project_name = project_name.strip()
        if clean_project_name and clean_project_name not in {"報告", "診斷報告", "診斷輔導報告"}:
            return clean_project_name
        return (normalized_document.hospital_name or clean_project_name or "report").strip()
